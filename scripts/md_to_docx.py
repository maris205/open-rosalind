#!/usr/bin/env python3
"""Convert a small Markdown file to DOCX.

This intentionally supports a conservative Markdown subset used by
Open-Rosalind Edu exports: headings, paragraphs, bullet lists, numbered lists,
fenced code blocks, and simple pipe tables.
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path


def require_docx():
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise SystemExit(
            "Missing dependency: python-docx. Install it with "
            "`pip install -r requirements.txt`."
        ) from exc
    return Document, Pt


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(set(cell.replace(" ", "")) <= {"-", ":"} for cell in cells)


def add_table(document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=0, cols=width)
    table.style = "Table Grid"
    for row in rows:
        cells = table.add_row().cells
        for index in range(width):
            cells[index].text = row[index] if index < len(row) else ""


def flush_paragraph(document, paragraph_lines: list[str]) -> None:
    if paragraph_lines:
        document.add_paragraph(" ".join(line.strip() for line in paragraph_lines).strip())
        paragraph_lines.clear()


def flush_code(document, code_lines: list[str], code_size) -> None:
    if code_lines:
        paragraph = document.add_paragraph()
        run = paragraph.add_run("\n".join(code_lines))
        run.font.name = "Consolas"
        run.font.size = code_size
        code_lines.clear()


def flush_table(document, table_rows: list[list[str]]) -> None:
    if table_rows:
        add_table(document, table_rows)
        table_rows.clear()


def convert_markdown_to_docx(input_path: Path, output_path: Path) -> None:
    Document, Pt = require_docx()
    document = Document()
    code_size = Pt(9)

    lines = input_path.read_text(encoding="utf-8").splitlines()
    paragraph_lines: list[str] = []
    code_lines: list[str] = []
    table_rows: list[list[str]] = []
    in_code = False

    for raw_line in lines:
        line = raw_line.rstrip()

        if line.startswith("```"):
            flush_paragraph(document, paragraph_lines)
            flush_table(document, table_rows)
            if in_code:
                flush_code(document, code_lines, code_size)
            in_code = not in_code
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            flush_paragraph(document, paragraph_lines)
            flush_table(document, table_rows)
            continue

        if line.startswith("|") and line.endswith("|"):
            flush_paragraph(document, paragraph_lines)
            if not is_table_separator(line):
                table_rows.append(split_table_row(line))
            continue
        flush_table(document, table_rows)

        if line.startswith("#"):
            flush_paragraph(document, paragraph_lines)
            level = min(len(line) - len(line.lstrip("#")), 4)
            title = line[level:].strip()
            document.add_heading(title, level=level)
            continue

        stripped = line.strip()
        if stripped.startswith(("- ", "* ")):
            flush_paragraph(document, paragraph_lines)
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue

        number_prefix = stripped.split(" ", 1)[0]
        if number_prefix.endswith(".") and number_prefix[:-1].isdigit() and " " in stripped:
            flush_paragraph(document, paragraph_lines)
            document.add_paragraph(stripped.split(" ", 1)[1].strip(), style="List Number")
            continue

        paragraph_lines.append(line)

    flush_paragraph(document, paragraph_lines)
    flush_code(document, code_lines, code_size)
    flush_table(document, table_rows)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Convert Markdown to DOCX.")
    parser.add_argument("input", type=Path, help="Input Markdown file")
    parser.add_argument("output", type=Path, help="Output DOCX file")
    args = parser.parse_args(argv)

    if not args.input.exists():
        print(f"Input file does not exist: {args.input}", file=sys.stderr)
        return 2

    convert_markdown_to_docx(args.input, args.output)
    print(f"Wrote {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
