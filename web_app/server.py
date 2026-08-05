#!/usr/bin/env python3
"""Local web UI for Open-Rosalind Edu.

The server is intentionally small and framework-free. It serves the static UI,
offers document upload text extraction, and wraps one OpenAI-compatible
chat-completions call for local use.
"""

from __future__ import annotations

import argparse
import email
import email.policy
import hashlib
import io
import json
import mimetypes
import os
import re
import shutil
import subprocess
import tarfile
import urllib.error
import urllib.request
import uuid
import sys
import threading
import time
from datetime import datetime, timezone
from http.cookies import SimpleCookie
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import quote, quote_plus, unquote, urlparse

if __package__:
    from .database import (
        authenticate_user, add_project_memory, approve_task_plan, claim_next_task_step, create_job as record_job_start,
        create_login_session, create_project, create_task_plan, create_user, delete_login_session,
        finish_job as record_job_finish, finish_task_step, get_project_workspace, get_task_plan, get_user_for_token,
        initialize_database, list_projects, project_memory_context, record_message_feedback, recover_interrupted_task_steps, retry_task_step, save_step_output_to_memory,
        user_owns_job, user_owns_project,
    )
    from .biology_tools import execute as execute_biology_tool
    from .openhands_runtime import (
        execute_step as execute_openhands_step,
        health as openhands_health,
        workspace_for_project as openhands_workspace_for_project,
    )
    from .task_queue import enqueue_plan_task, get_queue_job, queue_health
else:
    from database import (
        authenticate_user, add_project_memory, approve_task_plan, claim_next_task_step, create_job as record_job_start,
        create_login_session, create_project, create_task_plan, create_user, delete_login_session,
        finish_job as record_job_finish, finish_task_step, get_project_workspace, get_task_plan, get_user_for_token,
        initialize_database, list_projects, project_memory_context, record_message_feedback, recover_interrupted_task_steps, retry_task_step, save_step_output_to_memory,
        user_owns_job, user_owns_project,
    )
    from biology_tools import execute as execute_biology_tool
    from openhands_runtime import (
        execute_step as execute_openhands_step,
        health as openhands_health,
        workspace_for_project as openhands_workspace_for_project,
    )
    from task_queue import enqueue_plan_task, get_queue_job, queue_health


ROOT = Path(__file__).resolve().parents[1]
STATIC_DIR = Path(__file__).resolve().parent / "static"
SKILLS_DIR = ROOT / ".openhands" / "skills"
PROMPTS_DIR = ROOT / "prompts"
DEFAULT_BASE_URL = "https://llm-jl24o09ebj303z4e.cn-beijing.maas.aliyuncs.com/compatible-mode/v1"
DEFAULT_MODEL = "qwen3.7-max"
DESKTOP_MODE = os.environ.get("ROSALIND_DESKTOP_MODE", "0") == "1"
MAX_UPLOAD_BYTES = 12 * 1024 * 1024
MAX_EXTRACTED_CHARS = 120_000
MAX_CODE_CHARS = 50_000
MAX_LOG_CHARS = 100_000
MAX_OUTPUT_FILES = 30
MAX_OUTPUT_BYTES = 20 * 1024 * 1024
MAX_PROJECT_ARTIFACT_FILES = 200
MAX_PROJECT_ARTIFACT_BYTES = 100 * 1024 * 1024
EXECUTION_TIMEOUT_SECONDS = int(os.environ.get("ROSALIND_EXECUTION_TIMEOUT", "60"))
EXECUTION_MEMORY = os.environ.get("ROSALIND_EXECUTION_MEMORY", "512m")
EXECUTION_CPUS = os.environ.get("ROSALIND_EXECUTION_CPUS", "1")
EXECUTION_IMAGE = os.environ.get("ROSALIND_PYTHON_IMAGE", "python:3.12-slim")
JOBS_DIR = Path(os.environ.get("ROSALIND_JOBS_DIR", ROOT / "exports" / "jobs"))
EXECUTION_ENABLED = os.environ.get("ROSALIND_EXECUTION_ENABLED", "0") == "1"
NATIVE_EXECUTION_ENABLED = DESKTOP_MODE and os.environ.get("ROSALIND_DESKTOP_NATIVE_EXECUTION", "1") == "1"
EXECUTION_SLOTS = threading.BoundedSemaphore(int(os.environ.get("ROSALIND_EXECUTION_CONCURRENCY", "2")))
COOKIE_NAME = "rosalind_session"
COOKIE_SECURE = os.environ.get("ROSALIND_COOKIE_SECURE", "0") == "1"
AUTH_ATTEMPT_LIMIT = 10
AUTH_ATTEMPT_WINDOW = 15 * 60
AUTH_ATTEMPTS: dict[str, list[float]] = {}
AUTH_ATTEMPTS_LOCK = threading.Lock()
TASK_EXECUTION_LOCK = threading.Lock()
DESKTOP_PLAN_CONFIG_LOCK = threading.Lock()
DESKTOP_PLAN_MODEL_CONFIG: dict[str, dict[str, object]] = {}
TASK_SKILLS = {"agent-planner", "memory-manager", "evidence-manager", "reference-verification", "python-sandbox", "tool-audit", "agent-report-builder"}
DISCLAIMER = (
    "提示：当前为 Open-Rosalind Edu 模式，输出仅用于学习、写作辅助和初稿生成，"
    "不保证结论、引用、实验方案或统计解释完全正确。正式提交、发表或用于科研决策前，"
    "请人工核验原始文献、数据和引用。本系统不提供临床诊断或治疗建议。"
)


CONTENT_TYPES = {
    ".html": "text/html; charset=utf-8",
    ".css": "text/css; charset=utf-8",
    ".js": "application/javascript; charset=utf-8",
    ".json": "application/json; charset=utf-8",
    ".svg": "image/svg+xml",
}
TEXT_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".csv", ".tsv", ".json", ".xml", ".html", ".htm", ".rst",
    ".fa", ".fasta", ".faa", ".fna", ".ffn", ".frn", ".fastq", ".fq", ".gb", ".gbk",
    ".genbank", ".aln", ".clustal", ".phy", ".phylip", ".pdb", ".cif", ".mmcif",
}
SEQUENCE_EXTENSIONS = {
    ".fa", ".fasta", ".faa", ".fna", ".ffn", ".frn", ".fastq", ".fq", ".gb", ".gbk",
    ".genbank", ".aln", ".clustal", ".phy", ".phylip", ".pdb", ".cif", ".mmcif",
}
DOI_RE = re.compile(r"\b10\.\d{4,9}/[-._;()/:A-Z0-9]+", re.IGNORECASE)
PMID_RE = re.compile(r"\bPMID[:\s]*(\d{6,9})\b|\bPubMed[:\s]*(\d{6,9})\b", re.IGNORECASE)
YEAR_RE = re.compile(r"\b(19\d{2}|20\d{2})\b")
MARKDOWN_LINK_RE = re.compile(r"\[([^\]]+)\]\((https?://[^\s)]+)\)")
URL_RE = re.compile(r"https?://[^\s<>\]\)]+")


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8") if path.exists() else ""


def source_provider(url: str) -> str:
    host = urlparse(url).netloc.lower().removeprefix("www.")
    providers = {
        "pubmed.ncbi.nlm.nih.gov": "PubMed",
        "ncbi.nlm.nih.gov": "NCBI",
        "blast.ncbi.nlm.nih.gov": "NCBI BLAST",
        "rest.uniprot.org": "UniProt",
        "uniprot.org": "UniProt",
        "doi.org": "DOI",
        "crossref.org": "Crossref",
    }
    return next((label for domain, label in providers.items() if host == domain or host.endswith(f".{domain}")), host or "外部来源")


def extract_response_sources(content: str, limit: int = 30) -> list[dict[str, str]]:
    sources: list[dict[str, str]] = []
    seen: set[str] = set()

    def add(url: str, title: str = "") -> None:
        cleaned = url.rstrip(".,;:，。；：)")
        if not cleaned or cleaned in seen or len(sources) >= limit:
            return
        seen.add(cleaned)
        provider = source_provider(cleaned)
        sources.append({"title": title.strip() or provider, "url": cleaned, "provider": provider})

    for match in MARKDOWN_LINK_RE.finditer(content):
        add(match.group(2), match.group(1))
    for match in URL_RE.finditer(content):
        add(match.group(0))
    for doi in DOI_RE.findall(content):
        add(f"https://doi.org/{doi.rstrip('.,;:')}", f"DOI {doi.rstrip('.,;:')}")
    for match in PMID_RE.finditer(content):
        pmid = next((group for group in match.groups() if group), "")
        if pmid:
            add(f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/", f"PubMed PMID {pmid}")
    return sources


def response_trace(skill: str, mode: str, sources: list[dict[str, str]]) -> list[dict[str, object]]:
    trace: list[dict[str, object]] = [
        {
            "title": "识别任务与输出约束",
            "kind": "reasoning",
            "confidence": 72,
            "detail": "识别用户目标、所选 Skill、输入材料和需要明确标注的证据边界。",
        },
        {
            "title": f"应用 {skill} 专业流程",
            "kind": "reasoning",
            "confidence": 68,
            "detail": "按该 Skill 的结构组织分析；此步骤主要依赖模型推理，置信度低于真实工具查询。",
        },
    ]
    if sources:
        trace.append(
            {
                "title": "整理回答中的外部来源",
                "kind": "evidence",
                "confidence": 80,
                "detail": f"识别到 {len(sources)} 个可打开的外部链接；链接存在不代表其一定支持全部结论。",
            }
        )
    trace.append(
        {
            "title": "生成回答并标注证据边界",
            "kind": "reasoning" if mode == "llm" else "system",
            "confidence": 70 if mode == "llm" else 85,
            "detail": "输出为可审计的过程摘要，不包含模型内部隐藏思维链；科研结论仍需人工核验。",
        }
    )
    return trace


def biology_presentation(result: dict[str, object], content: str) -> tuple[list[dict[str, str]], list[dict[str, object]]]:
    sources = extract_response_sources(content)
    trace: list[dict[str, object]] = []
    if result.get("kind") == "protein":
        sequence = result.get("sequence") or {}
        entry = result.get("entry") or {}
        blast = result.get("blast") or {}
        if sequence:
            trace.append({"title": "解析并校验 FASTA", "kind": "tool", "confidence": 99, "detail": "本地确定性解析序列、长度、组成和异常字符。"})
        if entry:
            accession = str(entry.get("accession") or "")
            url = f"https://rest.uniprot.org/uniprotkb/{quote(accession)}.json" if accession else ""
            if url and not any(item["url"] == url for item in sources):
                sources.append({"title": f"UniProt {accession}", "url": url, "provider": "UniProt"})
            trace.append({"title": "查询 UniProt 注释", "kind": "tool", "confidence": 97, "detail": f"通过 UniProt REST API 获取 {accession or '匹配条目'} 的实时注释。"})
        if blast:
            rid = str(blast.get("rid") or "")
            if rid:
                blast_url = f"https://blast.ncbi.nlm.nih.gov/Blast.cgi?CMD=Get&RID={quote(rid)}"
                if not any(item["url"] == blast_url for item in sources):
                    sources.append({"title": f"NCBI BLAST RID {rid}", "url": blast_url, "provider": "NCBI BLAST"})
            confidence = 98 if blast.get("status") == "READY" else 82
            trace.append({"title": "运行 NCBI BLAST", "kind": "tool", "confidence": confidence, "detail": f"真实提交 blastp / swissprot；状态为 {blast.get('status', '未知')}。"})
        trace.append({"title": "汇总身份与功能证据", "kind": "reasoning", "confidence": 76, "detail": "基于工具结果组织解释；功能外推的置信度低于序列统计和数据库返回值。"})
    else:
        protein = result.get("protein") or {}
        diff = result.get("diff") or {}
        literature = result.get("literature") or []
        accession = str(protein.get("accession") or "")
        if accession:
            url = f"https://rest.uniprot.org/uniprotkb/{quote(accession)}.json"
            sources.append({"title": f"UniProt {accession}", "url": url, "provider": "UniProt"})
            trace.append({"title": "获取参考蛋白序列", "kind": "tool", "confidence": 97, "detail": f"通过 UniProt REST API 获取 {accession}。"})
        if diff:
            trace.append({"title": "计算 WT/MT 序列差异", "kind": "tool", "confidence": 99, "detail": "确定性比较位置、残基变化和长度变化。"})
        for item in literature:
            pmid = str(item.get("pmid") or "")
            if pmid:
                url = f"https://pubmed.ncbi.nlm.nih.gov/{quote(pmid)}/"
                sources.append({"title": str(item.get("title") or f"PMID {pmid}"), "url": url, "provider": "PubMed"})
        if literature:
            trace.append({"title": "检索 PubMed 文献", "kind": "tool", "confidence": 96, "detail": f"通过 NCBI E-utilities 获取 {len(literature)} 条实时书目信息。"})
        trace.append({"title": "评估分子影响与证据边界", "kind": "reasoning", "confidence": 62, "detail": "理化性质和结构影响属于启发式推断，不能替代 ClinVar/ACMG 临床分级。"})
    unique = {item["url"]: item for item in sources if item.get("url")}
    return list(unique.values())[:30], trace


def reference_presentation(result: dict[str, object]) -> tuple[list[dict[str, str]], list[dict[str, object]]]:
    sources: list[dict[str, str]] = []
    for item in result.get("results", []):
        metadata = item.get("metadata") or {}
        url = str(metadata.get("url") or "")
        if not url and item.get("doi"):
            url = f"https://doi.org/{item['doi']}"
        if not url and item.get("pmid"):
            url = f"https://pubmed.ncbi.nlm.nih.gov/{item['pmid']}/"
        if url:
            sources.append(
                {
                    "title": str(metadata.get("title") or item.get("reference") or "参考文献")[:300],
                    "url": url,
                    "provider": str(item.get("source") or source_provider(url)),
                }
            )
    unique = {item["url"]: item for item in sources}
    trace = [
        {"title": "解析参考文献标识符", "kind": "tool", "confidence": 98, "detail": "确定性提取 DOI、PMID、年份和书目文本。"},
        {"title": "查询 Crossref / PubMed", "kind": "tool", "confidence": 96, "detail": f"对 {result.get('count', 0)} 条记录执行实时元数据核验。"},
        {"title": "比较题名与年份一致性", "kind": "tool", "confidence": 90, "detail": "使用元数据匹配与题名重叠度标记不一致和疑似风险。"},
        {"title": "形成核验建议", "kind": "reasoning", "confidence": 78, "detail": "风险等级用于筛查，最终引用判断仍需打开原文人工确认。"},
    ]
    return list(unique.values())[:80], trace


def auth_rate_limited(address: str) -> bool:
    now = time.time()
    with AUTH_ATTEMPTS_LOCK:
        recent = [stamp for stamp in AUTH_ATTEMPTS.get(address, []) if now - stamp < AUTH_ATTEMPT_WINDOW]
        AUTH_ATTEMPTS[address] = recent
        return len(recent) >= AUTH_ATTEMPT_LIMIT


def record_auth_failure(address: str) -> None:
    with AUTH_ATTEMPTS_LOCK:
        AUTH_ATTEMPTS.setdefault(address, []).append(time.time())


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def execution_config() -> dict[str, object]:
    if NATIVE_EXECUTION_ENABLED:
        return {
            "enabled": True,
            "runtime": "Local Python",
            "image": "host-python",
            "network": "host",
            "readOnlyRoot": False,
            "cpu": "host",
            "memory": "host",
            "timeoutSeconds": EXECUTION_TIMEOUT_SECONDS,
            "maxCodeChars": MAX_CODE_CHARS,
            "maxOutputBytes": MAX_OUTPUT_BYTES,
        }
    return {
        "enabled": EXECUTION_ENABLED,
        "runtime": "Docker / Python",
        "image": EXECUTION_IMAGE,
        "network": "disabled",
        "readOnlyRoot": True,
        "cpu": EXECUTION_CPUS,
        "memory": EXECUTION_MEMORY,
        "timeoutSeconds": EXECUTION_TIMEOUT_SECONDS,
        "maxCodeChars": MAX_CODE_CHARS,
        "maxOutputBytes": MAX_OUTPUT_BYTES,
    }


def collect_output_files(output_dir: Path, job_id: str) -> tuple[list[dict[str, object]], list[str]]:
    files: list[dict[str, object]] = []
    warnings: list[str] = []
    total_bytes = 0
    for path in sorted(output_dir.rglob("*")):
        if path.is_symlink():
            warnings.append(f"忽略符号链接输出：{path.relative_to(output_dir).as_posix()}。")
            continue
        if not path.is_file():
            continue
        resolved = path.resolve()
        if not resolved.is_relative_to(output_dir.resolve()):
            warnings.append(f"忽略任务目录之外的输出：{path.name}。")
            continue
        relative = path.relative_to(output_dir).as_posix()
        size = path.stat().st_size
        total_bytes += size
        if len(files) >= MAX_OUTPUT_FILES:
            warnings.append(f"输出文件超过 {MAX_OUTPUT_FILES} 个，其余文件未列出。")
            break
        if total_bytes > MAX_OUTPUT_BYTES:
            warnings.append(f"输出文件总量超过 {MAX_OUTPUT_BYTES // 1024 // 1024} MB，超出部分未列出。")
            break
        files.append(
            {
                "name": relative,
                "size": size,
                "sha256": sha256_file(path),
                "url": f"/api/jobs/{job_id}/files/{quote(relative)}",
            }
        )
    return files, warnings


def collect_project_artifacts(workspace: Path, project_id: str) -> list[dict[str, object]]:
    """List user-visible Agent files while excluding runtime metadata and links."""
    root = workspace.resolve()
    artifacts: list[dict[str, object]] = []
    if not root.is_dir():
        return artifacts
    for path in root.rglob("*"):
        relative_parts = path.relative_to(root).parts
        if any(part.startswith(".") for part in relative_parts):
            continue
        if path.is_symlink() or not path.is_file():
            continue
        resolved = path.resolve()
        if not resolved.is_relative_to(root):
            continue
        stat = path.stat()
        if stat.st_size > MAX_PROJECT_ARTIFACT_BYTES:
            continue
        relative = path.relative_to(root).as_posix()
        artifacts.append(
            {
                "name": path.name,
                "path": relative,
                "size": stat.st_size,
                "mime": mimetypes.guess_type(path.name)[0] or "application/octet-stream",
                "modifiedAt": datetime.fromtimestamp(stat.st_mtime, timezone.utc).isoformat(),
                "url": f"/api/projects/{project_id}/artifacts/{quote(relative)}",
            }
        )
    artifacts.sort(key=lambda item: (str(item["modifiedAt"]), str(item["path"])), reverse=True)
    return artifacts[:MAX_PROJECT_ARTIFACT_FILES]


def export_container_outputs(container_name: str, output_dir: Path) -> str:
    export_code = (
        'import sys, tarfile; '
        'archive=tarfile.open(fileobj=sys.stdout.buffer, mode="w|"); '
        'archive.add("/workspace/output", arcname="."); archive.close()'
    )
    exported = subprocess.run(
        ["docker", "exec", container_name, "python", "-c", export_code],
        capture_output=True,
        timeout=20,
        check=False,
    )
    if exported.returncode != 0:
        return exported.stderr.decode("utf-8", errors="replace").strip() or "Container output export failed."
    total_size = 0
    root = output_dir.resolve()
    try:
        with tarfile.open(fileobj=io.BytesIO(exported.stdout), mode="r:*") as archive:
            for member in archive.getmembers():
                if member.isdir():
                    continue
                if not member.isfile() or member.issym() or member.islnk():
                    continue
                relative = Path(member.name)
                target = (root / relative).resolve()
                if not target.is_relative_to(root):
                    return f"Rejected unsafe output path: {member.name}"
                total_size += member.size
                if total_size > MAX_OUTPUT_BYTES:
                    return f"Output exceeds {MAX_OUTPUT_BYTES // 1024 // 1024} MB."
                target.parent.mkdir(parents=True, exist_ok=True)
                source = archive.extractfile(member)
                if source is None:
                    continue
                with target.open("wb") as destination:
                    shutil.copyfileobj(source, destination)
    except (tarfile.TarError, OSError) as exc:
        return f"Invalid container output archive: {exc}"
    return ""


def run_python_sandbox(code: str, user_id: str) -> dict[str, object]:
    if not EXECUTION_ENABLED and not NATIVE_EXECUTION_ENABLED:
        return {"ok": False, "error": "服务器未启用代码执行功能。"}
    if not isinstance(code, str) or not code.strip():
        return {"ok": False, "error": "Python 代码不能为空。"}
    if len(code) > MAX_CODE_CHARS:
        return {"ok": False, "error": f"代码过长，最多支持 {MAX_CODE_CHARS} 个字符。"}
    if NATIVE_EXECUTION_ENABLED:
        return run_native_python(code, user_id)
    if not shutil.which("docker"):
        return {"ok": False, "error": "服务器未安装 Docker。"}
    if not EXECUTION_SLOTS.acquire(blocking=False):
        return {"ok": False, "error": "执行队列已满，请稍后重试。", "busy": True}

    job_id = uuid.uuid4().hex
    container_name = f"rosalind-python-{job_id[:12]}"
    job_dir = JOBS_DIR / job_id
    input_dir = job_dir / "input"
    output_dir = job_dir / "output"
    input_dir.mkdir(parents=True, exist_ok=False)
    output_dir.mkdir(parents=True, exist_ok=False)
    job_dir.chmod(0o755)
    input_dir.chmod(0o755)
    script_path = input_dir / "main.py"
    script_path.write_text(code, encoding="utf-8")
    script_path.chmod(0o444)
    output_dir.chmod(0o755)

    create_command = [
        "docker", "create", "--name", container_name,
        "--network", "none",
        "--read-only",
        "--memory", EXECUTION_MEMORY,
        "--memory-swap", EXECUTION_MEMORY,
        "--cpus", EXECUTION_CPUS,
        "--pids-limit", "64",
        "--cap-drop", "ALL",
        "--security-opt", "no-new-privileges",
        "--user", "65534:65534",
        "--tmpfs", "/tmp:rw,noexec,nosuid,size=64m",
        "--tmpfs", f"/workspace/output:rw,noexec,nosuid,size={MAX_OUTPUT_BYTES},uid=65534,gid=65534,mode=0770",
        "--mount", f"type=bind,src={input_dir},dst=/workspace/input,readonly",
        "--workdir", "/workspace/output",
        EXECUTION_IMAGE,
        "/bin/sh", "-c",
        'python -I -B /workspace/input/main.py; code=$?; printf "%s" "$code" > /tmp/rosalind-exit-code; while :; do sleep 3600; done',
    ]
    started_at = time.time()
    code_sha256 = hashlib.sha256(code.encode("utf-8")).hexdigest()
    execution_limits = {
        "network": "disabled",
        "readOnlyRoot": True,
        "cpu": EXECUTION_CPUS,
        "memory": EXECUTION_MEMORY,
        "timeoutSeconds": EXECUTION_TIMEOUT_SECONDS,
        "maxOutputBytes": MAX_OUTPUT_BYTES,
    }
    try:
        record_job_start(job_id, user_id, code_sha256, EXECUTION_IMAGE, execution_limits, started_at)
    except Exception as exc:  # noqa: BLE001 - do not execute without a durable task record
        EXECUTION_SLOTS.release()
        return {"ok": False, "error": f"无法创建任务记录：{exc}"}
    status = "failed"
    exit_code: int | None = None
    stdout = ""
    stderr = ""
    container_created = False
    try:
        created = subprocess.run(create_command, capture_output=True, text=True, timeout=30, check=False)
        if created.returncode != 0:
            raise RuntimeError(created.stderr.strip() or "Docker container creation failed.")
        container_created = True
        started = subprocess.run(["docker", "start", container_name], capture_output=True, text=True, timeout=15, check=False)
        if started.returncode != 0:
            raise RuntimeError(started.stderr.strip() or "Docker container start failed.")
        deadline = time.monotonic() + EXECUTION_TIMEOUT_SECONDS
        while time.monotonic() < deadline:
            marker = subprocess.run(
                ["docker", "exec", container_name, "cat", "/tmp/rosalind-exit-code"],
                capture_output=True,
                text=True,
                timeout=5,
                check=False,
            )
            if marker.returncode == 0 and marker.stdout.strip():
                exit_code = int(marker.stdout.strip())
                break
            time.sleep(0.1)
        if exit_code is None:
            raise subprocess.TimeoutExpired(cmd="python sandbox", timeout=EXECUTION_TIMEOUT_SECONDS)
        logs = subprocess.run(["docker", "logs", container_name], capture_output=True, text=True, timeout=10, check=False)
        stdout = logs.stdout[-MAX_LOG_CHARS:]
        stderr = logs.stderr[-MAX_LOG_CHARS:]
        status = "succeeded" if exit_code == 0 else "failed"
    except subprocess.TimeoutExpired as exc:
        status = "timed_out"
        stdout = (exc.stdout or "")[-MAX_LOG_CHARS:] if isinstance(exc.stdout, str) else ""
        stderr = (exc.stderr or "")[-MAX_LOG_CHARS:] if isinstance(exc.stderr, str) else ""
        stderr = f"{stderr}\nExecution exceeded {EXECUTION_TIMEOUT_SECONDS} seconds.".strip()
    except Exception as exc:  # noqa: BLE001 - return a bounded execution error
        stderr = str(exc)
    finally:
        if container_created:
            if status != "timed_out":
                export_error = export_container_outputs(container_name, output_dir)
                if export_error:
                    stderr = f"{stderr}\nOutput collection failed: {export_error}".strip()
            subprocess.run(["docker", "rm", "-f", container_name], capture_output=True, timeout=10, check=False)
        EXECUTION_SLOTS.release()

    ended_at = time.time()
    files, warnings = collect_output_files(output_dir, job_id)
    audit = {
        "jobId": job_id,
        "status": status,
        "permissionLevel": 3,
        "codeSha256": code_sha256,
        "image": EXECUTION_IMAGE,
        "network": "disabled",
        "readOnlyRoot": True,
        "cpu": EXECUTION_CPUS,
        "memory": EXECUTION_MEMORY,
        "timeoutSeconds": EXECUTION_TIMEOUT_SECONDS,
        "startedAt": started_at,
        "endedAt": ended_at,
        "durationSeconds": round(ended_at - started_at, 3),
        "exitCode": exit_code,
        "outputs": files,
        "warnings": warnings,
    }
    (job_dir / "audit.json").write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    try:
        record_job_finish(job_id, audit, stdout, stderr, files)
    except Exception as exc:  # noqa: BLE001 - preserve file audit even if DB update fails
        audit["warnings"].append(f"数据库任务收尾失败：{exc}")
    return {
        "ok": status == "succeeded",
        "jobId": job_id,
        "status": status,
        "stdout": stdout,
        "stderr": stderr,
        "files": files,
        "audit": audit,
    }


def run_native_python(code: str, user_id: str) -> dict[str, object]:
    if not EXECUTION_SLOTS.acquire(blocking=False):
        return {"ok": False, "error": "执行队列已满，请稍后重试。", "busy": True}

    job_id = uuid.uuid4().hex
    job_dir = JOBS_DIR / job_id
    input_dir = job_dir / "input"
    output_dir = job_dir / "output"
    input_dir.mkdir(parents=True, exist_ok=False)
    output_dir.mkdir(parents=True, exist_ok=False)
    script_path = input_dir / "main.py"
    script_path.write_text(code, encoding="utf-8")
    code_sha256 = hashlib.sha256(code.encode("utf-8")).hexdigest()
    started_at = time.time()
    execution_limits = {
        "network": "host",
        "readOnlyRoot": False,
        "cpu": "host",
        "memory": "host",
        "timeoutSeconds": EXECUTION_TIMEOUT_SECONDS,
        "maxOutputBytes": MAX_OUTPUT_BYTES,
    }
    try:
        record_job_start(job_id, user_id, code_sha256, "host-python", execution_limits, started_at)
    except Exception as exc:  # noqa: BLE001 - do not execute without a durable task record
        EXECUTION_SLOTS.release()
        return {"ok": False, "error": f"无法创建任务记录：{exc}"}

    status = "failed"
    exit_code: int | None = None
    stdout = ""
    stderr = ""
    try:
        completed = subprocess.run(
            [sys.executable, "-I", "-B", str(script_path)],
            cwd=output_dir,
            capture_output=True,
            text=True,
            timeout=EXECUTION_TIMEOUT_SECONDS,
            check=False,
            env={
                **os.environ,
                "HOME": str(job_dir),
                "TMPDIR": str(job_dir),
                "ROSALIND_OUTPUT_DIR": str(output_dir),
            },
        )
        exit_code = completed.returncode
        stdout = completed.stdout[-MAX_LOG_CHARS:]
        stderr = completed.stderr[-MAX_LOG_CHARS:]
        status = "succeeded" if exit_code == 0 else "failed"
    except subprocess.TimeoutExpired as exc:
        status = "timed_out"
        stdout = (exc.stdout or "")[-MAX_LOG_CHARS:] if isinstance(exc.stdout, str) else ""
        stderr = (exc.stderr or "")[-MAX_LOG_CHARS:] if isinstance(exc.stderr, str) else ""
        stderr = f"{stderr}\nExecution exceeded {EXECUTION_TIMEOUT_SECONDS} seconds.".strip()
    except Exception as exc:  # noqa: BLE001 - return a bounded execution error
        stderr = str(exc)
    finally:
        EXECUTION_SLOTS.release()

    ended_at = time.time()
    files, warnings = collect_output_files(output_dir, job_id)
    warnings.insert(0, "代码在本机 Python 中运行，不具备 Docker 文件系统或网络隔离。")
    audit = {
        "jobId": job_id,
        "status": status,
        "permissionLevel": 3,
        "codeSha256": code_sha256,
        "image": "host-python",
        "network": "host",
        "readOnlyRoot": False,
        "cpu": "host",
        "memory": "host",
        "timeoutSeconds": EXECUTION_TIMEOUT_SECONDS,
        "startedAt": started_at,
        "endedAt": ended_at,
        "durationSeconds": round(ended_at - started_at, 3),
        "exitCode": exit_code,
        "outputs": files,
        "warnings": warnings,
    }
    (job_dir / "audit.json").write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    try:
        record_job_finish(job_id, audit, stdout, stderr, files)
    except Exception as exc:  # noqa: BLE001 - preserve file audit even if DB update fails
        audit["warnings"].append(f"数据库任务收尾失败：{exc}")
    return {
        "ok": status == "succeeded",
        "jobId": job_id,
        "status": status,
        "stdout": stdout,
        "stderr": stderr,
        "files": files,
        "audit": audit,
    }


def decode_text_file(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")



def strip_bibtex_value(value: str) -> str:
    value = value.strip().rstrip(",").strip()
    if (value.startswith("{") and value.endswith("}")) or (value.startswith('"') and value.endswith('"')):
        value = value[1:-1]
    value = re.sub(r"[{}]", "", value)
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def split_bibtex_entries(text: str) -> list[str]:
    entries: list[str] = []
    index = 0
    while True:
        start = text.find("@", index)
        if start == -1:
            break
        brace = text.find("{", start)
        if brace == -1:
            break
        depth = 0
        end = brace
        while end < len(text):
            char = text[end]
            if char == "{":
                depth += 1
            elif char == "}":
                depth -= 1
                if depth == 0:
                    entries.append(text[start : end + 1])
                    index = end + 1
                    break
            end += 1
        else:
            break
    return entries


def parse_bibtex_fields(entry: str) -> dict[str, str]:
    brace = entry.find("{")
    if brace == -1:
        return {}
    body = entry[brace + 1 : -1].strip()
    comma = body.find(",")
    if comma == -1:
        return {}
    body = body[comma + 1 :]
    fields: dict[str, str] = {}
    index = 0
    while index < len(body):
        match = re.search(r"([A-Za-z][A-Za-z0-9_-]*)\s*=\s*", body[index:])
        if not match:
            break
        key = match.group(1).lower()
        value_start = index + match.end()
        if value_start >= len(body):
            break
        if body[value_start] == "{":
            depth = 0
            pos = value_start
            while pos < len(body):
                if body[pos] == "{":
                    depth += 1
                elif body[pos] == "}":
                    depth -= 1
                    if depth == 0:
                        pos += 1
                        break
                pos += 1
            raw_value = body[value_start:pos]
            index = pos + 1
        elif body[value_start] == '"':
            pos = value_start + 1
            while pos < len(body):
                if body[pos] == '"' and body[pos - 1] != "\\":
                    pos += 1
                    break
                pos += 1
            raw_value = body[value_start:pos]
            index = pos + 1
        else:
            pos = body.find(",", value_start)
            if pos == -1:
                pos = len(body)
            raw_value = body[value_start:pos]
            index = pos + 1
        fields[key] = strip_bibtex_value(raw_value)
    return fields


def format_bibtex_reference(fields: dict[str, str]) -> str:
    authors = fields.get("author", "").replace(" and ", "; ")
    title = fields.get("title", "")
    venue = fields.get("journal") or fields.get("journaltitle") or fields.get("booktitle") or fields.get("publisher", "")
    year = fields.get("year", "")
    volume = fields.get("volume", "")
    number = fields.get("number", "")
    pages = fields.get("pages", "")
    doi = fields.get("doi", "")
    pmid = fields.get("pmid", "")
    parts = [part for part in [authors, title, venue, year] if part]
    ref = ". ".join(parts)
    details: list[str] = []
    if volume:
        details.append(f"vol. {volume}")
    if number:
        details.append(f"no. {number}")
    if pages:
        details.append(f"pages {pages}")
    if details:
        ref = f"{ref}. {', '.join(details)}"
    if doi:
        ref = f"{ref}. doi:{doi}"
    if pmid:
        ref = f"{ref}. PMID:{pmid}"
    return ref.strip().rstrip(".") + "."


def extract_bibtex(data: bytes) -> str:
    text = decode_text_file(data)
    references: list[str] = []
    for entry in split_bibtex_entries(text):
        fields = parse_bibtex_fields(entry)
        reference = format_bibtex_reference(fields)
        if reference.strip() != ".":
            references.append(reference)
    if not references:
        raise ValueError("未能从 BibTeX 中解析到参考文献条目。")
    return "\n".join(references)
def extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("当前环境缺少 python-docx，无法解析 DOCX。") from exc

    from io import BytesIO

    document = Document(BytesIO(data))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ValueError("当前环境缺少 pypdf，无法解析 PDF。") from exc

    from io import BytesIO

    reader = PdfReader(BytesIO(data))
    parts: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            parts.append(f"[Page {index}]\n{text}")
    return "\n\n".join(parts)


def truncate_extracted_text(text: str) -> tuple[str, bool]:
    if len(text) <= MAX_EXTRACTED_CHARS:
        return text, False
    return text[:MAX_EXTRACTED_CHARS], True


def extract_uploaded_document(filename: str, data: bytes) -> dict[str, object]:
    suffix = Path(filename).suffix.lower()
    kind = "document"
    if suffix == ".bib":
        text = extract_bibtex(data)
        kind = "bibliography"
    elif suffix in TEXT_EXTENSIONS:
        text = decode_text_file(data)
        if suffix in SEQUENCE_EXTENSIONS:
            kind = "sequence"
    elif suffix == ".docx":
        text = extract_docx(data)
    elif suffix == ".pdf":
        text = extract_pdf(data)
        kind = "paper"
    else:
        raise ValueError(
            f"暂不支持的文件类型：{suffix or 'unknown'}。支持 PDF、BibTeX、DOCX、文本、"
            "FASTA、FASTQ、GenBank、PDB 和 mmCIF。"
        )

    text, truncated = truncate_extracted_text(text)
    if not text.strip():
        raise ValueError("未能从文件中提取到可用文本。扫描版 PDF 需要先 OCR。")
    return {
        "filename": filename,
        "size": len(data),
        "extension": suffix,
        "kind": kind,
        "text": text,
        "chars": len(text),
        "truncated": truncated,
    }
def parse_multipart_file(headers, body: bytes) -> tuple[str, bytes]:
    content_type = headers.get("Content-Type", "")
    if "multipart/form-data" not in content_type:
        raise ValueError("上传请求必须使用 multipart/form-data。")

    message = email.message_from_bytes(
        b"Content-Type: " + content_type.encode("utf-8") + b"\r\n\r\n" + body,
        policy=email.policy.default,
    )
    for part in message.iter_parts():
        if part.get_content_disposition() != "form-data":
            continue
        if part.get_param("name", header="content-disposition") != "document":
            continue
        filename = part.get_filename() or "uploaded.txt"
        data = part.get_payload(decode=True) or b""
        return filename, data
    raise ValueError("没有找到名为 document 的上传文件。")



def clean_identifier(value: str) -> str:
    return value.strip().rstrip(".,;)]}")


def extract_doi(reference: str) -> str:
    match = DOI_RE.search(reference)
    return clean_identifier(match.group(0)) if match else ""


def extract_pmid(reference: str) -> str:
    match = PMID_RE.search(reference)
    if not match:
        return ""
    return next(group for group in match.groups() if group)


def extract_year(reference: str) -> str:
    match = YEAR_RE.search(reference)
    return match.group(1) if match else ""


def normalize_tokens(text: str) -> set[str]:
    text = re.sub(r"[^\w\s]", " ", text.lower())
    stopwords = {"the", "and", "for", "with", "from", "this", "that", "into", "using", "study", "analysis", "of", "in", "on", "a", "an", "to"}
    return {token for token in text.split() if len(token) > 2 and token not in stopwords}


def title_overlap(title: str, reference: str) -> float:
    title_tokens = normalize_tokens(title)
    if not title_tokens:
        return 0.0
    reference_tokens = normalize_tokens(reference)
    return len(title_tokens & reference_tokens) / len(title_tokens)


def parse_reference_lines(text: str) -> list[str]:
    if re.search(r"@\w+\s*{", text):
        try:
            return [line for line in extract_bibtex(text.encode("utf-8")).splitlines() if line.strip()]
        except ValueError:
            pass
    lines: list[str] = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        line = re.sub(r"^\s*(\[\d+\]|\d+[.)]|[-*•])\s*", "", line).strip()
        if line:
            lines.append(line)
    if len(lines) <= 1 and ";" in text:
        candidates = [part.strip() for part in text.split(";") if len(part.strip()) > 30]
        return candidates or lines
    return lines
def fetch_json(url: str, timeout: int = 20) -> tuple[dict | None, str]:
    request = urllib.request.Request(url, headers={"User-Agent": "Open-Rosalind-Edu/0.2 (local reference verifier)"})
    try:
        with urllib.request.urlopen(request, timeout=timeout) as response:
            return json.loads(response.read().decode("utf-8")), ""
    except urllib.error.HTTPError as exc:
        return None, f"HTTP {exc.code}"
    except Exception as exc:  # noqa: BLE001 - report lookup failures to UI
        return None, str(exc)


def crossref_by_doi(doi: str) -> tuple[dict | None, str]:
    data, error = fetch_json(f"https://api.crossref.org/works/{quote(doi, safe='')}")
    if not data:
        return None, error
    return data.get("message"), ""


def crossref_search(reference: str) -> tuple[dict | None, str]:
    data, error = fetch_json(f"https://api.crossref.org/works?rows=3&query.bibliographic={quote_plus(reference)}")
    if not data:
        return None, error
    items = data.get("message", {}).get("items", [])
    return (items[0] if items else None), "" if items else "No Crossref candidates"


def pubmed_by_pmid(pmid: str) -> tuple[dict | None, str]:
    url = f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi?db=pubmed&id={quote_plus(pmid)}&retmode=json"
    data, error = fetch_json(url)
    if not data:
        return None, error
    result = data.get("result", {})
    return result.get(pmid), "" if result.get(pmid) else "PMID not found"


def crossref_metadata(item: dict | None) -> dict[str, str]:
    if not item:
        return {}
    authors = item.get("author") or []
    first_author = ""
    if authors:
        first = authors[0]
        first_author = " ".join(part for part in [first.get("given", ""), first.get("family", "")] if part).strip()
    year = ""
    date_parts = (item.get("issued") or {}).get("date-parts") or []
    if date_parts and date_parts[0]:
        year = str(date_parts[0][0])
    return {
        "title": (item.get("title") or [""])[0],
        "journal": (item.get("container-title") or [""])[0],
        "year": year,
        "doi": item.get("DOI", ""),
        "firstAuthor": first_author,
        "url": item.get("URL", ""),
    }


def pubmed_metadata(item: dict | None) -> dict[str, str]:
    if not item:
        return {}
    authors = item.get("authors") or []
    first_author = authors[0].get("name", "") if authors else ""
    year_match = YEAR_RE.search(item.get("pubdate", ""))
    return {
        "title": item.get("title", ""),
        "journal": item.get("fulljournalname", "") or item.get("source", ""),
        "year": year_match.group(1) if year_match else "",
        "doi": "",
        "firstAuthor": first_author,
        "url": f"https://pubmed.ncbi.nlm.nih.gov/{item.get('uid', '')}/" if item.get("uid") else "",
    }


def assess_reference(reference: str) -> dict[str, object]:
    doi = extract_doi(reference)
    pmid = extract_pmid(reference)
    stated_year = extract_year(reference)
    checks: list[str] = []
    warnings: list[str] = []
    metadata: dict[str, str] = {}
    source = ""
    label = "Unverified"

    if doi:
        item, error = crossref_by_doi(doi)
        if item:
            metadata = crossref_metadata(item)
            source = "Crossref DOI"
            label = "Verified"
            checks.append("DOI exists in Crossref.")
        else:
            warnings.append(f"DOI lookup failed: {error or 'not found'}")
            label = "Fabrication Risk"
    elif pmid:
        item, error = pubmed_by_pmid(pmid)
        if item:
            metadata = pubmed_metadata(item)
            source = "PubMed PMID"
            label = "Verified"
            checks.append("PMID exists in PubMed.")
        else:
            warnings.append(f"PMID lookup failed: {error or 'not found'}")
            label = "Fabrication Risk"
    else:
        item, error = crossref_search(reference)
        if item:
            metadata = crossref_metadata(item)
            source = "Crossref bibliographic search"
            overlap = title_overlap(metadata.get("title", ""), reference)
            checks.append(f"Best Crossref candidate title overlap: {overlap:.2f}.")
            label = "Candidate Match" if overlap >= 0.45 else "Unverified"
            if overlap < 0.45:
                warnings.append("No DOI/PMID supplied and best title match is weak.")
        else:
            warnings.append(error or "No bibliographic candidate found.")

    if metadata:
        if stated_year and metadata.get("year") and stated_year != metadata["year"]:
            label = "Metadata Mismatch"
            warnings.append(f"Year mismatch: reference says {stated_year}, source says {metadata['year']}.")
        if metadata.get("title"):
            overlap = title_overlap(metadata["title"], reference)
            if overlap and overlap < 0.35:
                label = "Metadata Mismatch" if label == "Verified" else label
                warnings.append(f"Title overlap is low ({overlap:.2f}); manually compare the title.")

    if label in {"Unverified", "Fabrication Risk"}:
        warnings.append("Do not use this reference until manually verified from DOI, PMID, publisher page, or library database.")

    return {
        "reference": reference,
        "label": label,
        "doi": doi,
        "pmid": pmid,
        "statedYear": stated_year,
        "source": source,
        "metadata": metadata,
        "checks": checks,
        "warnings": warnings,
    }


def verify_references_text(text: str) -> dict[str, object]:
    references = parse_reference_lines(text)
    if not references:
        return {"ok": False, "error": "没有识别到参考文献条目。请每条参考文献单独成行。"}
    results = [assess_reference(reference) for reference in references[:80]]
    counts: dict[str, int] = {}
    for item in results:
        label = str(item["label"])
        counts[label] = counts.get(label, 0) + 1
    return {
        "ok": True,
        "count": len(results),
        "truncated": len(references) > len(results),
        "counts": counts,
        "results": results,
    }


def reference_report_markdown(result: dict[str, object]) -> str:
    if not result.get("ok"):
        return str(result.get("error", "Reference verification failed."))
    lines = ["# 参考文献验证报告", "", "## 总览", ""]
    counts = result.get("counts", {})
    for label in ["Verified", "Metadata Mismatch", "Candidate Match", "Unverified", "Fabrication Risk"]:
        lines.append(f"- {label}: {counts.get(label, 0)}")
    if result.get("truncated"):
        lines.append("- 注意：条目超过 80 条，本次只验证前 80 条。")
    lines.extend(["", "## 逐条结果", ""])
    for index, item in enumerate(result.get("results", []), start=1):
        metadata = item.get("metadata") or {}
        lines.append(f"### {index}. {item.get('label')}")
        lines.append("")
        lines.append(f"原始条目：{item.get('reference')}")
        if item.get("doi"):
            lines.append(f"DOI：{item.get('doi')}")
        if item.get("pmid"):
            lines.append(f"PMID：{item.get('pmid')}")
        if item.get("source"):
            lines.append(f"核验来源：{item.get('source')}")
        if metadata:
            lines.append(f"匹配题名：{metadata.get('title', '')}")
            lines.append(f"期刊：{metadata.get('journal', '')}")
            lines.append(f"年份：{metadata.get('year', '')}")
            lines.append(f"第一作者：{metadata.get('firstAuthor', '')}")
            if metadata.get("url"):
                lines.append(f"链接：{metadata.get('url')}")
        warnings = item.get("warnings") or []
        if warnings:
            lines.append("风险提示：")
            for warning in warnings:
                lines.append(f"- {warning}")
        lines.append("")
    lines.extend([
        "## 人工核验清单",
        "",
        "- 对 Metadata Mismatch、Unverified 和 Fabrication Risk 条目逐条打开 DOI、PMID、出版社页面或图书馆数据库核验。",
        "- 不要仅凭题名相似就认定引用真实存在。",
        "- 参考文献真实存在不等于支持正文 claim；高风险医学 claim 还需要阅读全文核验。",
        "",
        DISCLAIMER,
    ])
    return "\n".join(lines)
def summarize_skill(skill_dir: Path) -> dict[str, str]:
    text = read_text(skill_dir / "SKILL.md")
    title = skill_dir.name.replace("_", " ").title()
    for line in text.splitlines():
        if line.startswith("# "):
            title = line[2:].strip()
            break
    purpose = ""
    lines = text.splitlines()
    for index, line in enumerate(lines):
        if line.strip() == "## Purpose" and index + 2 < len(lines):
            purpose = lines[index + 2].strip()
            break
    return {
        "id": skill_dir.name,
        "title": title,
        "purpose": purpose,
        "content": text,
    }


def load_skills() -> list[dict[str, str]]:
    if not SKILLS_DIR.exists():
        return []
    return [summarize_skill(path) for path in sorted(SKILLS_DIR.iterdir()) if path.is_dir()]


def build_messages(skill: str, user_input: str, history: list[dict[str, str]] | None = None) -> list[dict[str, str]]:
    if not re.fullmatch(r"[a-z0-9_-]+", skill) or not (SKILLS_DIR / skill / "SKILL.md").is_file():
        raise ValueError("不支持的 Agent Skill。")
    system_prompt = read_text(PROMPTS_DIR / "system_edu.md")
    writing_policy = read_text(PROMPTS_DIR / "writing_policy.md")
    citation_policy = read_text(PROMPTS_DIR / "citation_policy.md")
    skill_prompt = read_text(SKILLS_DIR / skill / "SKILL.md")
    system = "\n\n".join(part for part in [system_prompt, writing_policy, citation_policy, skill_prompt] if part)
    messages = [{"role": "system", "content": system}]
    for item in (history or [])[-8:]:
        role = item.get("role")
        content = str(item.get("content", "")).strip()
        if role in {"user", "assistant"} and content:
            messages.append({"role": role, "content": content[:8000]})
    user = (
        "请根据所选 Open-Rosalind Edu Skill 处理以下输入。"
        "默认用中文、Markdown 输出，并在结尾附上 Edu Mode disclaimer。\n\n"
        f"{user_input.strip()}"
    )
    messages.append({"role": "user", "content": user})
    return messages


def call_openai_compatible(payload: dict) -> dict:
    api_key = payload.get("apiKey") or os.environ.get("DASHSCOPE_API_KEY", "")
    if not api_key:
        return {
            "ok": False,
            "mode": "prompt_only",
            "error": "未提供 API Key。可在页面临时填写，或启动服务前设置 DASHSCOPE_API_KEY。",
            "prompt": payload.get("promptPreview", ""),
        }

    base_url = (payload.get("baseUrl") or os.environ.get("QWEN_BASE_URL") or DEFAULT_BASE_URL).rstrip("/")
    model = payload.get("model") or os.environ.get("QWEN_MODEL") or DEFAULT_MODEL
    body = {
        "model": model,
        "messages": payload["messages"],
        "temperature": float(payload.get("temperature", 0.3)),
    }
    request = urllib.request.Request(
        f"{base_url}/chat/completions",
        data=json.dumps(body, ensure_ascii=False).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=120) as response:
            data = json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        if exc.code in {401, 403}:
            error = (
                "模型服务认证失败。请清空设置中的临时 API Key 后重试；"
                "如果仍然失败，请联系管理员检查服务器模型配置。"
            )
        elif exc.code == 429:
            error = "模型服务当前请求过多或额度不足，请稍后重试。"
        else:
            error = f"模型服务请求失败（HTTP {exc.code}），请稍后重试。"
        return {"ok": False, "mode": "error", "error": error, "status": exc.code}
    except Exception as exc:  # noqa: BLE001 - surface local web errors to UI
        return {"ok": False, "mode": "error", "error": str(exc)}

    content = data.get("choices", [{}])[0].get("message", {}).get("content", "")
    return {"ok": True, "mode": "llm", "content": content, "raw": data}


def parse_json_object(text: str) -> dict:
    fenced = re.search(r"```(?:json)?\s*([\s\S]*?)```", text, re.IGNORECASE)
    candidate = fenced.group(1).strip() if fenced else text.strip()
    start = candidate.find("{")
    end = candidate.rfind("}")
    if start == -1 or end <= start:
        raise ValueError("模型没有返回有效的 JSON 任务计划。")
    return json.loads(candidate[start : end + 1])


def request_model_config(payload: dict[str, object]) -> dict[str, object]:
    return {
        name: payload[name]
        for name in ("apiKey", "baseUrl", "model")
        if isinstance(payload.get(name), str) and str(payload[name]).strip()
    }


def generate_task_plan_with_model(
    goal: str,
    memory: list[dict[str, object]],
    model_config: dict[str, object] | None = None,
) -> tuple[list[dict[str, str]], str]:
    memory_text = "\n".join(
        f"- [{item['category']}] {str(item['content'])[:2000]}" for item in memory
    ) or "- No stored project memory."
    system = (
        "You are the Open-Rosalind biomedical research task planner. Return JSON only. "
        "Do not fabricate evidence, citations, datasets, results, or tool execution. "
        "Create a short reviewable plan whose steps can be executed independently by existing skills."
    )
    user = f"""Create a plan for this research goal:

{goal}

Project memory:
{memory_text}

Return exactly this JSON shape with 2-6 steps:
{{
  "summary": "short plan summary",
  "steps": [
    {{
      "title": "step title",
      "instruction": "self-contained instruction including expected output and evidence boundary",
      "skill": "one allowed skill"
    }}
  ]
}}

Allowed skills: {', '.join(sorted(TASK_SKILLS))}.
Use python-sandbox only to prepare reviewable offline Python code; never claim code has executed."""
    result = call_openai_compatible(
        {
            **(model_config or {}),
            "messages": [{"role": "system", "content": system}, {"role": "user", "content": user}],
            "temperature": 0.2,
        }
    )
    if not result.get("ok"):
        raise RuntimeError(str(result.get("error", "任务计划生成失败。")))
    content = str(result.get("content", ""))
    payload = parse_json_object(content)
    raw_steps = payload.get("steps")
    if not isinstance(raw_steps, list) or not 1 <= len(raw_steps) <= 10:
        raise ValueError("模型返回的任务步骤数量无效。")
    steps: list[dict[str, str]] = []
    for index, raw in enumerate(raw_steps, start=1):
        if not isinstance(raw, dict):
            raise ValueError(f"任务步骤 {index} 格式无效。")
        skill = str(raw.get("skill", "agent-planner"))
        if skill not in TASK_SKILLS:
            skill = "agent-planner"
        steps.append(
            {
                "title": str(raw.get("title", f"步骤 {index}")),
                "instruction": str(raw.get("instruction", "")),
                "skill": skill,
            }
        )
    return steps, str(payload.get("summary", ""))


def execute_claimed_task_step(user_id: str, plan_data: dict[str, object], step_data: dict[str, object]) -> dict[str, object]:
    if not step_data:
        return get_task_plan(user_id, str(plan_data["id"])) or {}
    memory = project_memory_context(user_id, str(plan_data["projectId"]), limit=30)
    memory_text = "\n".join(
        f"- [{item['category']}] {str(item['content'])[:2000]}" for item in memory
    ) or "- No stored memory."
    previous_text = "\n\n".join(
        f"Previous step {item['position']} - {item['title']}:\n{str(item['output'])[:8000]}"
        for item in plan_data.get("previous", [])
    ) or "No previous completed steps."
    task_input = f"""Project goal:
{plan_data['goal']}

Current step:
{step_data['instruction']}

Project memory:
{memory_text}

Previous completed outputs:
{previous_text}

Complete only the current step. Clearly distinguish evidence, inference, and unverified content. Do not claim any tool or code ran unless a real tool log is included above."""
    messages = build_messages(str(step_data.get("skill", "agent-planner")), task_input, history=[])
    runtime = os.environ.get("ROSALIND_AGENT_RUNTIME", "legacy")
    if runtime == "openhands":
        try:
            system_prompt = (
                f"{messages[0]['content']}\n\n"
                "You are executing inside the OpenHands agent runtime with terminal and filesystem tools. "
                "Use the dedicated project workspace for any files you create. The selected Rosalind Skill "
                "instructions are included above. "
                "Use tools only when the task requires them, record what actually ran, and never claim an "
                "external query, computation, or file operation occurred unless you executed it. Finish with "
                "a self-contained Markdown result for the current task step."
            )
            workspace = openhands_workspace_for_project(user_id, str(plan_data["projectId"]))
            openhands = execute_openhands_step(
                system_prompt,
                messages[-1]["content"],
                workspace_path=workspace,
            )
            marker = f"<!-- runtime: openhands; conversation: {openhands.get('conversationId', '')} -->\n"
            return finish_task_step(user_id, str(step_data["id"]), output=marker + str(openhands["content"])) or {}
        except Exception as exc:  # noqa: BLE001 - persist runtime failure or use explicit fallback
            if os.environ.get("OPENHANDS_FALLBACK_ENABLED", "1") != "1":
                return finish_task_step(user_id, str(step_data["id"]), error=f"OpenHands 执行失败：{exc}") or {}
    with DESKTOP_PLAN_CONFIG_LOCK:
        model_config = dict(DESKTOP_PLAN_MODEL_CONFIG.get(str(plan_data["id"])) or {})
    result = call_openai_compatible({**model_config, "messages": messages, "temperature": 0.2})
    if result.get("ok"):
        return finish_task_step(user_id, str(step_data["id"]), output=str(result.get("content", ""))) or {}
    return finish_task_step(user_id, str(step_data["id"]), error=str(result.get("error", "步骤执行失败。"))) or {}


def run_next_task_step(user_id: str, plan_id: str) -> dict[str, object]:
    with TASK_EXECUTION_LOCK:
        claimed = claim_next_task_step(user_id, plan_id)
        if not claimed:
            raise PermissionError("任务计划不存在。")
        try:
            return execute_claimed_task_step(user_id, *claimed)
        except Exception as exc:  # noqa: BLE001 - persist a retryable failed step
            step_data = claimed[1]
            if step_data:
                return finish_task_step(user_id, str(step_data["id"]), error=f"后台步骤执行失败：{exc}") or {}
            raise


def run_all_task_steps(user_id: str, plan_id: str) -> dict[str, object]:
    plan: dict[str, object] = {}
    for _ in range(10):
        plan = run_next_task_step(user_id, plan_id)
        if plan.get("status") in {"completed", "failed"}:
            if DESKTOP_MODE:
                with DESKTOP_PLAN_CONFIG_LOCK:
                    DESKTOP_PLAN_MODEL_CONFIG.pop(plan_id, None)
            return plan
    return plan


class Handler(BaseHTTPRequestHandler):
    server_version = "OpenRosalindEdu/0.2"

    def log_message(self, format: str, *args) -> None:  # noqa: A002
        print(f"{self.address_string()} - {format % args}")

    def send_json(self, data: dict | list, status: int = 200, headers: dict[str, str] | None = None) -> None:
        raw = json.dumps(data, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(raw)))
        for name, value in (headers or {}).items():
            self.send_header(name, value)
        self.end_headers()
        self.wfile.write(raw)

    def read_json(self) -> dict:
        length = int(self.headers.get("Content-Length", "0"))
        raw = self.rfile.read(length).decode("utf-8") if length else "{}"
        return json.loads(raw)

    def read_body(self) -> bytes:
        length = int(self.headers.get("Content-Length", "0"))
        if length > MAX_UPLOAD_BYTES:
            raise ValueError(f"文件过大，最大支持 {MAX_UPLOAD_BYTES // 1024 // 1024} MB。")
        return self.rfile.read(length)

    def session_token(self) -> str:
        cookie = SimpleCookie()
        cookie.load(self.headers.get("Cookie", ""))
        morsel = cookie.get(COOKIE_NAME)
        return morsel.value if morsel else ""

    def current_user(self) -> dict[str, str] | None:
        return get_user_for_token(self.session_token())

    def require_user(self) -> dict[str, str] | None:
        user = self.current_user()
        if not user:
            self.send_json({"ok": False, "error": "请先登录。"}, status=401)
        return user

    def auth_cookie(self, token: str, max_age: int) -> str:
        parts = [f"{COOKIE_NAME}={token}", "Path=/", "HttpOnly", "SameSite=Lax", f"Max-Age={max_age}"]
        if COOKIE_SECURE:
            parts.append("Secure")
        return "; ".join(parts)

    def do_GET(self) -> None:  # noqa: N802
        if self.path == "/api/auth/me":
            user = self.current_user()
            if not user:
                self.send_json({"ok": False, "authenticated": False}, status=401)
            else:
                self.send_json({"ok": True, "authenticated": True, "user": user})
            return
        if self.path == "/api/config":
            self.send_json(
                {
                    "baseUrl": os.environ.get("QWEN_BASE_URL", DEFAULT_BASE_URL),
                    "model": os.environ.get("QWEN_MODEL", DEFAULT_MODEL),
                    "hasEnvApiKey": bool(os.environ.get("DASHSCOPE_API_KEY")),
                    "disclaimer": DISCLAIMER,
                    "desktopMode": DESKTOP_MODE,
                }
            )
            return
        if self.path == "/api/desktop/status":
            if not DESKTOP_MODE:
                self.send_error(404)
                return
            self.send_json(
                {
                    "ok": True,
                    "desktopMode": True,
                    "execution": "local",
                    "queue": "in-process",
                    "python": sys.version.split()[0],
                    "dockerAvailable": bool(shutil.which("docker")),
                    "agentRuntime": os.environ.get("ROSALIND_AGENT_RUNTIME", "legacy"),
                }
            )
            return
        if self.path == "/api/execution/config":
            self.send_json(execution_config())
            return
        if self.path == "/api/runtime/status":
            user = self.require_user()
            if not user:
                return
            runtime = os.environ.get("ROSALIND_AGENT_RUNTIME", "legacy")
            details = openhands_health() if runtime == "openhands" else {"ok": True, "runtime": "legacy"}
            self.send_json({"ok": bool(details.get("ok")), "runtime": runtime, "details": details}, status=200 if details.get("ok") else 503)
            return
        if self.path == "/api/projects":
            user = self.require_user()
            if user:
                self.send_json({"ok": True, "projects": list_projects(user["id"])})
            return
        if self.path == "/api/queue/status":
            user = self.require_user()
            if not user:
                return
            try:
                self.send_json(queue_health())
            except Exception as exc:  # noqa: BLE001
                self.send_json({"ok": False, "error": f"任务队列不可用：{exc}"}, status=503)
            return
        task_status_match = re.fullmatch(r"/api/tasks/([a-f0-9]{32})/status", self.path)
        if task_status_match:
            user = self.require_user()
            if not user:
                return
            try:
                task = get_queue_job(user["id"], task_status_match.group(1))
            except Exception as exc:  # noqa: BLE001
                self.send_json({"ok": False, "error": f"任务队列不可用：{exc}"}, status=503)
                return
            if not task:
                self.send_error(404)
            else:
                self.send_json({"ok": True, "task": task})
            return
        workspace_match = re.fullmatch(r"/api/projects/([a-f0-9-]{36})/workspace", self.path)
        if workspace_match:
            user = self.require_user()
            if not user:
                return
            workspace = get_project_workspace(user["id"], workspace_match.group(1))
            if not workspace:
                self.send_error(404)
            else:
                self.send_json({"ok": True, **workspace})
            return
        request_path = unquote(self.path.split("?", 1)[0])
        artifacts_match = re.fullmatch(r"/api/projects/([a-f0-9-]{36})/artifacts", request_path)
        if artifacts_match:
            user = self.require_user()
            if not user:
                return
            project_id = artifacts_match.group(1)
            if not user_owns_project(user["id"], project_id):
                self.send_error(404)
                return
            workspace = openhands_workspace_for_project(user["id"], project_id)
            self.send_json({"ok": True, "artifacts": collect_project_artifacts(workspace, project_id)})
            return
        artifact_download_match = re.fullmatch(r"/api/projects/([a-f0-9-]{36})/artifacts/(.+)", request_path)
        if artifact_download_match:
            user = self.require_user()
            if not user:
                return
            project_id, relative_name = artifact_download_match.groups()
            if not user_owns_project(user["id"], project_id):
                self.send_error(404)
                return
            relative_parts = Path(relative_name).parts
            if not relative_parts or any(part.startswith(".") for part in relative_parts):
                self.send_error(404)
                return
            workspace = openhands_workspace_for_project(user["id"], project_id).resolve()
            candidate = workspace.joinpath(*relative_parts)
            linked_path = any(workspace.joinpath(*relative_parts[:index]).is_symlink() for index in range(1, len(relative_parts) + 1))
            try:
                file_path = candidate.resolve()
                invalid_file = (
                    linked_path
                    or not file_path.is_relative_to(workspace)
                    or not file_path.is_file()
                    or file_path.stat().st_size > MAX_PROJECT_ARTIFACT_BYTES
                )
            except (OSError, ValueError):
                invalid_file = True
            if invalid_file:
                self.send_error(404)
                return
            raw = file_path.read_bytes()
            encoded_name = quote(file_path.name, safe="")
            self.send_response(200)
            self.send_header("Content-Type", mimetypes.guess_type(file_path.name)[0] or "application/octet-stream")
            self.send_header("Content-Disposition", f"attachment; filename=artifact; filename*=UTF-8''{encoded_name}")
            self.send_header("Content-Length", str(len(raw)))
            self.end_headers()
            self.wfile.write(raw)
            return
        if self.path.startswith("/api/jobs/") and "/files/" in self.path:
            user = self.require_user()
            if not user:
                return
            match = re.fullmatch(r"/api/jobs/([a-f0-9]{32})/files/(.+)", unquote(self.path.split("?", 1)[0]))
            if not match:
                self.send_error(404)
                return
            job_id, relative_name = match.groups()
            if not user_owns_job(user["id"], job_id):
                self.send_error(404)
                return
            output_root = (JOBS_DIR / job_id / "output").resolve()
            file_path = (output_root / relative_name).resolve()
            if not file_path.is_relative_to(output_root) or not file_path.is_file():
                self.send_error(404)
                return
            raw = file_path.read_bytes()
            self.send_response(200)
            self.send_header("Content-Type", mimetypes.guess_type(file_path.name)[0] or "application/octet-stream")
            self.send_header("Content-Disposition", f'attachment; filename="{file_path.name}"')
            self.send_header("Content-Length", str(len(raw)))
            self.end_headers()
            self.wfile.write(raw)
            return
        if self.path == "/api/skills":
            self.send_json(load_skills())
            return

        path = "/" if self.path == "/" else unquote(self.path.split("?", 1)[0])
        if path == "/":
            file_path = STATIC_DIR / ("index.html" if DESKTOP_MODE else "landing.html")
        elif path in {"/app", "/app/"}:
            file_path = STATIC_DIR / "index.html"
        else:
            file_path = STATIC_DIR / path.lstrip("/")
        if not file_path.exists() or not file_path.resolve().is_relative_to(STATIC_DIR.resolve()):
            self.send_error(404)
            return
        raw = file_path.read_bytes()
        self.send_response(200)
        self.send_header("Content-Type", CONTENT_TYPES.get(file_path.suffix, "application/octet-stream"))
        self.send_header("Content-Length", str(len(raw)))
        self.end_headers()
        self.wfile.write(raw)

    def do_POST(self) -> None:  # noqa: N802
        if self.path in {"/api/auth/register", "/api/auth/login"}:
            address = self.client_address[0]
            if auth_rate_limited(address):
                self.send_json({"ok": False, "error": "尝试次数过多，请稍后再试。"}, status=429)
                return
            try:
                payload = self.read_json()
                email_value = str(payload.get("email", ""))
                password_value = str(payload.get("password", ""))
                if self.path == "/api/auth/register":
                    user = create_user(email_value, password_value)
                else:
                    user = authenticate_user(email_value, password_value)
                    if not user:
                        record_auth_failure(address)
                        raise ValueError("邮箱或密码错误。")
                token, expires_at = create_login_session(user["id"])
                max_age = max(0, int((expires_at - datetime.now(timezone.utc)).total_seconds()))
                self.send_json({"ok": True, "user": user}, headers={"Set-Cookie": self.auth_cookie(token, max_age)})
            except (json.JSONDecodeError, ValueError) as exc:
                self.send_json({"ok": False, "error": str(exc)}, status=400)
            return
        if self.path == "/api/auth/logout":
            delete_login_session(self.session_token())
            self.send_json({"ok": True}, headers={"Set-Cookie": self.auth_cookie("", 0)})
            return

        if not self.require_user():
            return
        user = self.current_user()
        if self.path == "/api/messages/feedback":
            try:
                payload = self.read_json()
                feedback = record_message_feedback(
                    user["id"],
                    str(payload.get("messageId", "")),
                    str(payload.get("chatId", "")),
                    str(payload.get("skill", "")),
                    str(payload.get("rating", "")),
                    str(payload.get("content", ""))[:100_000],
                )
                self.send_json({"ok": True, "feedback": feedback})
            except (json.JSONDecodeError, ValueError) as exc:
                self.send_json({"ok": False, "error": str(exc)}, status=400)
            return
        if self.path == "/api/projects":
            try:
                payload = self.read_json()
                project = create_project(user["id"], str(payload.get("name", "")), str(payload.get("description", "")))
                self.send_json({"ok": True, "project": project}, status=201)
            except ValueError as exc:
                self.send_json({"ok": False, "error": str(exc)}, status=400)
            return
        memory_match = re.fullmatch(r"/api/projects/([a-f0-9-]{36})/memory", self.path)
        if memory_match:
            try:
                payload = self.read_json()
                memory = add_project_memory(memory_match.group(1), user["id"], str(payload.get("category", "fact")), str(payload.get("content", "")))
                self.send_json({"ok": True, "memory": memory})
            except (ValueError, PermissionError) as exc:
                self.send_json({"ok": False, "error": str(exc)}, status=400)
            return
        plan_generate_match = re.fullmatch(r"/api/projects/([a-f0-9-]{36})/plans/generate", self.path)
        if plan_generate_match:
            try:
                project_id = plan_generate_match.group(1)
                if not user_owns_project(user["id"], project_id):
                    self.send_error(404)
                    return
                payload = self.read_json()
                goal = str(payload.get("goal", "")).strip()
                if not goal:
                    raise ValueError("请输入任务目标。")
                model_config = request_model_config(payload)
                steps, summary = generate_task_plan_with_model(
                    goal,
                    project_memory_context(user["id"], project_id),
                    model_config,
                )
                plan = create_task_plan(project_id, user["id"], goal, steps)
                if DESKTOP_MODE:
                    with DESKTOP_PLAN_CONFIG_LOCK:
                        DESKTOP_PLAN_MODEL_CONFIG[str(plan["id"])] = model_config
                self.send_json({"ok": True, "plan": plan, "summary": summary})
            except (ValueError, RuntimeError) as exc:
                self.send_json({"ok": False, "error": str(exc)}, status=400)
            return
        plan_action_match = re.fullmatch(r"/api/plans/([a-f0-9-]{36})/(confirm|run-next|run-all)", self.path)
        if plan_action_match:
            plan_id, action = plan_action_match.groups()
            try:
                if action == "confirm":
                    plan = approve_task_plan(user["id"], plan_id)
                else:
                    plan = get_task_plan(user["id"], plan_id)
                    if not plan:
                        self.send_error(404)
                        return
                    if plan.get("status") not in {"approved", "running"}:
                        raise ValueError("计划尚未确认，或当前状态不能执行。")
                    task = enqueue_plan_task(user["id"], plan_id, "next" if action == "run-next" else "all")
                    self.send_json({"ok": True, "plan": plan, "task": task}, status=202)
                    return
                if not plan:
                    self.send_error(404)
                else:
                    self.send_json({"ok": True, "plan": plan})
            except (ValueError, PermissionError, RuntimeError) as exc:
                self.send_json({"ok": False, "error": str(exc)}, status=400)
            except Exception as exc:  # noqa: BLE001
                self.send_json({"ok": False, "error": f"后台任务提交失败：{exc}"}, status=503)
            return
        step_retry_match = re.fullmatch(r"/api/steps/([a-f0-9-]{36})/retry", self.path)
        if step_retry_match:
            try:
                plan = retry_task_step(user["id"], step_retry_match.group(1))
                self.send_json({"ok": True, "plan": plan}) if plan else self.send_error(404)
            except ValueError as exc:
                self.send_json({"ok": False, "error": str(exc)}, status=400)
            return
        step_memory_match = re.fullmatch(r"/api/steps/([a-f0-9-]{36})/memory", self.path)
        if step_memory_match:
            try:
                payload = self.read_json()
                memory = save_step_output_to_memory(user["id"], step_memory_match.group(1), str(payload.get("category", "conclusion")))
                self.send_json({"ok": True, "memory": memory})
            except (ValueError, PermissionError) as exc:
                self.send_json({"ok": False, "error": str(exc)}, status=400)
            return
        if self.path == "/api/execute/python":
            try:
                payload = self.read_json()
                if payload.get("confirmed") is not True:
                    raise ValueError("执行前必须由用户明确确认代码和权限。")
                self.send_json(run_python_sandbox(str(payload.get("code", "")), user["id"]))
            except ValueError as exc:
                self.send_json({"ok": False, "error": str(exc)}, status=400)
            except RuntimeError as exc:
                self.send_json({"ok": False, "error": str(exc)}, status=503)
            return
        if self.path == "/api/biology/analyze":
            try:
                payload = self.read_json()
                tool = str(payload.get("tool", ""))
                query = str(payload.get("input", ""))
                attachment = str(payload.get("attachment", ""))
                result = execute_biology_tool(tool, query, attachment)
                sources, trace = biology_presentation(result.get("result") or {}, str(result.get("content", "")))
                result["sources"] = sources
                result["trace"] = trace
                self.send_json(result)
            except (json.JSONDecodeError, ValueError) as exc:
                self.send_json({"ok": False, "error": str(exc)}, status=400)
            except Exception as exc:  # noqa: BLE001 - report bounded external-tool failures
                self.send_json({"ok": False, "error": f"生物学工具执行失败：{exc}"}, status=502)
            return
        if self.path == "/api/verify-references":
            try:
                payload = self.read_json()
                input_value = payload.get("input", "")
                if not isinstance(input_value, str):
                    input_value = json.dumps(input_value, ensure_ascii=False)
                result = verify_references_text(input_value)
                if result.get("ok"):
                    result["content"] = reference_report_markdown(result)
                    result["sources"], result["trace"] = reference_presentation(result)
                    self.send_json(result)
                else:
                    self.send_json(result, status=400)
            except Exception as exc:  # noqa: BLE001 - surface local verification errors to UI
                self.send_json({"ok": False, "error": f"参考文献验证失败：{exc}"}, status=500)
            return
        if self.path == "/api/upload":
            try:
                filename, data = parse_multipart_file(self.headers, self.read_body())
                result = extract_uploaded_document(filename, data)
                self.send_json({"ok": True, **result})
            except ValueError as exc:
                self.send_json({"ok": False, "error": str(exc)}, status=400)
            except Exception as exc:  # noqa: BLE001 - surface local parse errors to UI
                self.send_json({"ok": False, "error": f"解析失败：{exc}"}, status=500)
            return

        if self.path != "/api/generate":
            self.send_error(404)
            return
        try:
            payload = self.read_json()
            skill = str(payload.get("skill", "paper_summary"))
            user_input = payload.get("input", "")
            if not isinstance(user_input, str):
                raise ValueError("输入内容必须是文本。")
            history = payload.get("history") if isinstance(payload.get("history"), list) else []
            messages = build_messages(skill, user_input, history)
        except (json.JSONDecodeError, ValueError) as exc:
            self.send_json({"ok": False, "error": str(exc)}, status=400)
            return
        prompt_preview = "\n\n".join(f"{item['role'].upper()}:\n{item['content']}" for item in messages)
        payload["messages"] = messages
        payload["promptPreview"] = prompt_preview
        result = call_openai_compatible(payload)
        if not result.get("ok") and result.get("mode") == "prompt_only":
            result["content"] = (
                "## 待发送 Prompt\n\n"
                "当前未配置 API Key。你可以复制下面的 prompt 到任意兼容模型，"
                "或在左侧临时填写 API Key 后重新生成。\n\n"
                "```text\n"
                f"{prompt_preview}\n"
                "```\n\n"
                f"{DISCLAIMER}"
            )
        content = str(result.get("content", ""))
        sources = extract_response_sources(content)
        result["sources"] = sources
        result["trace"] = response_trace(skill, str(result.get("mode", "llm")), sources)
        self.send_json(result)


def main() -> int:
    parser = argparse.ArgumentParser(description="Run the local Open-Rosalind Edu web UI.")
    parser.add_argument("--host", default="127.0.0.1")
    parser.add_argument("--port", default=8765, type=int)
    args = parser.parse_args()
    if DESKTOP_MODE and args.host not in {"127.0.0.1", "localhost", "::1"}:
        parser.error("Desktop mode may only bind to a loopback address.")

    initialize_database()
    if DESKTOP_MODE:
        recovered = recover_interrupted_task_steps()
        if recovered:
            print(f"Recovered {len(recovered)} interrupted desktop plan(s) as failed.")
    server = ThreadingHTTPServer((args.host, args.port), Handler)
    print(f"Open-Rosalind Edu local web UI: http://{args.host}:{args.port}")
    server.serve_forever()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
